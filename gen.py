#!/usr/bin/env python3


from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os;

from docx.shared import Mm, Pt, Inches
from docx.oxml import OxmlElement
import lxml
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import yaml
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from datetime import datetime
date = datetime.now().strftime("%Y-%m-%d")
# refer to docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapSquare wrapText = "bothSides"/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )


# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename    
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

# refer to docx.oxml.shape.__init__.py
register_element_cls('wp:anchor', CT_Anchor)

font_name_light = 'Montserrat-Light'
font_name_strong = "Raleway[Semi-bold 600]"
# /usr/share/fonts/truetype/open-sans/OpenSans-CondLight.ttf
# /usr/share/fonts/truetype/open-sans/OpenSans-Light.ttf
# /usr/share/fonts/truetype/roboto/unhinted/RobotoTTF/Roboto-Light.ttf
# RobotoCondensed-Light
font_name = 'OpenSans-Light'#

def indent_table(table, indent):
    # noinspection PyProtectedMember
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
def set_text_on_cell(cell: _Cell, text: str, font_size: int = 11, paragraphs = [], line_height: float = 0.9, bold: bool = False, font_name=font_name, **kwargs):
    set_cell_border(
        cell,
        top={"sz": 0, "space": "0"},
        bottom={"sz": 0, "space": "0"},
        start={"sz": 0, "space": "0"},
        end={"sz": 0, "space": "0"},
    )
    if text:
        paragraph = cell.paragraphs[0]
        paragraph.text = text
        paragraph.line_spacing = line_height
        if bold:
            paragraph.runs[0].bold = bold
        font = paragraph.runs[0].font
        font.name = font_name
        font.size = Pt(font_size)
    else:
        delete_paragraph(cell.paragraphs[0])

    if len(paragraphs) > 0:
        for p in paragraphs:
            paragraph = cell.add_paragraph()
            paragraph.text = p['text']
            if 'bold' in p and p['bold']:
                paragraph.runs[0].bold = True
            paragraph.line_spacing = line_height
            if 'bullet' in p and p['bullet']:
                paragraph.text = f"- " + p['text']
                paragraph.style = list_style
            font = paragraph.runs[0].font
            font.name = font_name
            font.size = Pt(font_size)
    return paragraph

document = Document("template.docx")
styles = document.styles
bullet_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
bullet_style_paragraph_format = bullet_style.paragraph_format
bullet_style_paragraph_format.left_indent = Inches(0.25)
bullet_style_paragraph_format.first_line_indent = Inches(-0.09)
bullet_style_paragraph_format.space_before = Pt(2)
bullet_style_paragraph_format.widow_control = True

title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH, builtin=False)
title_style.font.name = font_name
title_style.font.size = Pt(18)

for style in styles:
    print(f"{style.name} -- {style.type}")
list_style = styles['List Bullet']

section = document.sections[0]

# section.header_distance = Mm(52)
section.left_margin = Mm(23)
section.right_margin = Mm(23)
# section.footer_distance = Mm(30)

header = section.header
# for table in header.tables:
# indent_table(header.tables[0], -Mm(23))
header.tables[0].style.paragraph_format.left_indent = -Mm(23)
header.tables[0].style.paragraph_format.right_indent = -Mm(23)
for paragraph in header.paragraphs:
    paragraph.paragraph_format.left_indent = -Mm(23)
    paragraph.paragraph_format.right_indent = -Mm(23)
    paragraph.style = document.styles["Header"]
    for run in paragraph.runs:
        run.font.name = font_name
# paragraph.text = 'TEXT'

def with_indent(pp):
    pp.paragraph_format.left_indent = Mm(3)
    pp.paragraph_format.right_indent = Mm(3)
    return pp

#header.paragraphs.pop()

def fill_table_with_values(table, values):
    for i, faq in enumerate(values):
        row = table.add_row()
        set_text_on_cell(row.cells[0], faq['title'], font_size=11, font_name=font_name_strong)
        if type(faq['values']) == str:
            set_text_on_cell( row.cells[1], faq['values'], font_size=9)
            continue
        text_paragraphs = []
        for value in faq['values']:
            text_paragraphs.append({
                'text': value,
                'bullet': True,
            })
        text_paragraphs.append({
            'text': '\n',
        })
        set_text_on_cell(row.cells[1], None, font_size=9, paragraphs=text_paragraphs)
# paragraph.text = "Left Text\tCenter Text\tRight Text"
def add_project(cells, project):
    title_paragraph = set_text_on_cell(cells[0], f"{project['start']} - {project['end']}", font_size=11)
    title_paragraph.runs[0].bold = True

    text_paragraphs = [
        {
            'text': project['client'] + " / " + project['industry'],
            'bold': True,
        },
    ]

    points_text = ""
    for point in project['key_points']:
        text_paragraphs.append({
            'bold': False,
            'text': point,
            'bullet': True,
        })
    
    text_paragraphs.append({
        'bold': False,
        'text': '',
        'bullet': False
    })
    set_text_on_cell(cells[1], project['role'] + " for", font_size=9, paragraphs=text_paragraphs)
    return cells
with open('projects.yaml') as f:
    data = yaml.load(f, Loader=yaml.FullLoader)

    pp = with_indent(document.add_paragraph())
    p_font = pp.add_run()

    p_font.add_text(data['introduction'])

    p_font.line_spacing = 0.9
    p_font.font.name = font_name
    p_font.font.size = Pt(9)
    p_font.add_break()
    p_font.add_break()
    p_font.add_break()
    p_font.add_break()

    profile_table = document.add_table(0, 2)
    # indent_table(profile_table, Mm(99))
    fill_table_with_values(profile_table, data['faq_sheet'])

    for row in profile_table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Mm(8)
    document.add_page_break()

    document.add_heading('IT-Knowledge', 0)
    p_font = document.add_paragraph().add_run("""
    +++ excellent Knowledge, acquired through years of project work
    ++ very good knowledge, acquired through at least one longer project assignment
    + basic knowledge, e.g. through prototypes or occasional use 
    """)
    p_font.line_spacing = 0.9
    p_font.font.name = font_name


    it_knowledge_table = document.add_table(0, 2)
    fill_table_with_values(it_knowledge_table, data['knowledge_graph'])


    for row in it_knowledge_table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Mm(15)





    document.add_page_break()
    document.add_heading('Project History', 0).runs[0].font.name = font_name
    projects_table = document.add_table(0, 2)
    
    for i, project in enumerate(data['history']):
        row = projects_table.add_row()
        add_project(row.cells, project)
        row1 = projects_table.add_row()
        set_text_on_cell(row1.cells[0], "Used Technologies", font_size=11, line_height = 1.1)
        set_text_on_cell(row1.cells[1], project['tech'], font_size=9, line_height = 1.1)
    document.add_page_break()
    
    document.add_heading('Permanent Positions', 0).runs[0].font.name = font_name
    pos_table = document.add_table(0, 2)
    for i, project in enumerate(data['permanent_positions']):
        row = pos_table.add_row()
        set_text_on_cell(row.cells[0], f"{project['start']} - {project['end']}", font_size=11, font_name=font_name_strong)
        text_paragraphs = []
        for point in project['key_points']:
            text_paragraphs.append({
                'bold': False,
                'text': point,
                'bullet': True,
            })

        set_text_on_cell(row.cells[1], f"Role:", font_size=9, paragraphs=text_paragraphs)
    document.add_heading('Open-Source Engagement', 0).runs[0].font.name = font_name
    links_table = document.add_table(0, 2)

    for i, project in enumerate(data['links']):
        row = links_table.add_row()
        paragraph = set_text_on_cell(row.cells[0], f"{project['title']}", font_size=11, font_name=font_name_strong)
        set_text_on_cell(row.cells[1], f"{project['link']}", font_size=9)

    

    fname = f'Florian-Kasper-{date}.docx'
    document.save(fname)
    os.system(f"UNOPATH=/usr/lib/libreoffice /usr/bin/python3 /usr/bin/unoconv -fpdf {fname}")