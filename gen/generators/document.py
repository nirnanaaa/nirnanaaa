from docx import Document as DocxDocument
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


ACCENT = RGBColor(242, 183, 5)
DARK = RGBColor(19, 22, 38)
MID_GRAY = RGBColor(102, 102, 102)
FONT = "Helvetica"

LABEL_WIDTH = Mm(50)
MARGIN = Mm(20)


def _no_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    ))


def _full_width_table(table):
    table._tbl.tblPr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    ))


def _run(p, text, size=10, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _sp(p, before=0, after=0):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _gold_border(p):
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="4" w:color="F2B705"/>'
        '</w:pBdr>'
    ))


def _setup_header(section):
    header = section.header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    # Name
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, before=0, after=0)
    _run(p, "Florian Kasper", size=20, bold=True, color=DARK)

    # Title
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, before=0, after=2)
    _run(p, "Software and DevOps Engineer", size=10, color=MID_GRAY)

    # Contact line
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, before=4, after=6)
    _run(p, "info@florian-kasper.com", size=8, color=MID_GRAY)
    _run(p, "   \u00b7   ", size=8, color=ACCENT)
    _run(p, "+49 1516 7522873", size=8, color=MID_GRAY)
    _run(p, "   \u00b7   ", size=8, color=ACCENT)
    _run(p, "florian-kasper.com", size=8, color=MID_GRAY)

    # Gold rule
    _gold_border(p)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    _sp(p, before=24, after=8)
    _run(p, text, size=13, bold=True, color=DARK)
    _gold_border(p)


def _two_col_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    _full_width_table(table)
    table.columns[0].width = LABEL_WIDTH

    for i, (label, values) in enumerate(rows_data):
        lc = table.cell(i, 0)
        _no_borders(lc)
        lp = lc.paragraphs[0]
        _sp(lp, before=8, after=8)
        _run(lp, label, size=10, bold=True)

        vc = table.cell(i, 1)
        _no_borders(vc)
        for j, v in enumerate(values):
            vp = vc.paragraphs[0] if j == 0 else vc.add_paragraph()
            _sp(vp, before=8, after=3)
            _run(vp, v, size=10)

    return table


def _separator(doc):
    p = doc.add_paragraph()
    _sp(p, before=6, after=6)
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="2" w:space="1" w:color="D0D0D0"/>'
        '</w:pBdr>'
    ))


class Document():
    def filename(self, lang):
        return f"cv_{lang}.docx"

    def _add_overview(self, doc, cv):
        p = doc.add_paragraph()
        _sp(p, before=6, after=16)
        _run(p, cv['introduction'].strip(), size=10)

        rows = [(item['title'], item['values']) for item in cv['faq_sheet']]
        _two_col_table(doc, rows)

    def _add_skills(self, doc, cv):
        _section_heading(doc, "IT-Knowledge")

        legend = [
            "+++ excellent knowledge, acquired through years of project work",
            "++ very good knowledge, acquired through at least one longer project assignment",
            "+ basic knowledge, e.g. through prototypes or occasional use",
        ]
        for line in legend:
            p = doc.add_paragraph()
            _sp(p, before=1, after=1)
            _run(p, line, size=7, bold=True, italic=True, color=MID_GRAY)

        sp = doc.add_paragraph()
        _sp(sp, before=6, after=0)

        rows = [(item['title'], item['values']) for item in cv['knowledge_graph']]
        _two_col_table(doc, rows)

    def _add_projects(self, doc, history, title="Project History"):
        _section_heading(doc, title)

        for idx, data in enumerate(history):
            if idx > 0:
                _separator(doc)

            table = doc.add_table(rows=1, cols=2)
            _full_width_table(table)
            table.columns[0].width = LABEL_WIDTH

            dc = table.cell(0, 0)
            _no_borders(dc)
            dp = dc.paragraphs[0]
            _sp(dp, before=10, after=0)
            _run(dp, f"{data['start']} -\n{data['end']}", size=9, color=MID_GRAY)

            rc = table.cell(0, 1)
            _no_borders(rc)

            rp = rc.paragraphs[0]
            _sp(rp, before=10, after=3)
            role_text = f"{data['role']} for {data['client']}"
            if data.get('industry'):
                role_text += f"  \u00b7  {data['industry']}"
            _run(rp, role_text, size=10, bold=True)

            if data.get('tech'):
                tp = rc.add_paragraph()
                _sp(tp, before=2, after=6)
                _run(tp, data['tech'], size=8, italic=True, color=MID_GRAY)

            for kp in data['key_points']:
                bp = rc.add_paragraph()
                _sp(bp, before=3, after=3)
                _run(bp, f"\u2022  {kp}", size=9)

            ep = rc.add_paragraph()
            _sp(ep, before=0, after=0)

    def _add_links(self, doc):
        _section_heading(doc, "Links")
        _two_col_table(doc, [("GitHub", ["https://github.com/nirnanaaa"])])

    def create(self, cv, output_file):
        doc = DocxDocument()

        section = doc.sections[0]
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header_distance = Mm(5)

        style = doc.styles['Normal']
        style.font.name = FONT
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        _setup_header(section)

        self._add_overview(doc, cv)
        doc.add_page_break()
        self._add_skills(doc, cv)
        doc.add_page_break()
        self._add_projects(doc, cv['history'])
        self._add_projects(doc, cv['permanent_positions'], "Employment")
        self._add_links(doc)

        doc.save(output_file)
